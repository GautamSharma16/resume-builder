#!/usr/bin/env python3
import json
import sys


def emit(text, engine):
    print(json.dumps({"text": text or "", "engine": engine}, ensure_ascii=False))


def extract_pdf(path):
    try:
        import pdfplumber

        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                parts = []
                text = page.extract_text(x_tolerance=1, y_tolerance=3, layout=True) or ""
                if text.strip():
                    parts.append(text)

                tables = page.extract_tables() or []
                for table in tables:
                    rows = []
                    for row in table or []:
                        cells = [(cell or "").strip() for cell in row or []]
                        if any(cells):
                            rows.append(" | ".join(cells))
                    if rows:
                        parts.append("\n".join(rows))

                if parts:
                    pages.append("\n".join(parts))

        text = "\n\n".join(pages).strip()
        if text:
            return text, "pdfplumber"
    except Exception:
        pass

    try:
        import fitz

        doc = fitz.open(path)
        pages = []
        for page in doc:
            text = page.get_text("text", sort=True) or ""
            if text.strip():
                pages.append(text)
        text = "\n\n".join(pages).strip()
        if text:
            return text, "pymupdf"
    except Exception:
        pass

    try:
        import fitz
        import pytesseract
        from PIL import Image

        doc = fitz.open(path)
        pages = []
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(image) or ""
            if text.strip():
                pages.append(text)
        text = "\n\n".join(pages).strip()
        if text:
            return text, "ocr"
    except Exception:
        pass

    return "", "none"


def extract_docx(path):
    try:
        from docx import Document

        document = Document(path)
        blocks = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            if "heading" in style:
                blocks.append(text.upper())
            elif paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
                blocks.append("- " + text)
            else:
                blocks.append(text)

        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                blocks.append("\n".join(rows))

        text = "\n\n".join(blocks).strip()
        if text:
            return text, "python-docx"
    except Exception:
        pass

    return "", "none"


def main():
    if len(sys.argv) < 3:
        emit("", "none")
        return

    kind = sys.argv[1].lower()
    path = sys.argv[2]

    if kind == "pdf":
        text, engine = extract_pdf(path)
    elif kind == "docx":
        text, engine = extract_docx(path)
    else:
        text, engine = "", "none"

    emit(text, engine)


if __name__ == "__main__":
    main()
